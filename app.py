from flask import Flask, request as flask_request, send_file, jsonify, redirect
from flask_cors import CORS
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, io, os, json, time
import requests as http_requests

app = Flask(__name__)
CORS(app, expose_headers=['X-Drive-File-Id', 'X-Job-Folder-Id'])

# ── CONFIG ──
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'EXTERIOR_MASTER_TEMPLATE.docx')
CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID', '568040559683-4jt2t8u4me7oimp1etb0evt5nknsgg34.apps.googleusercontent.com')
CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET', '')
SCOPES = 'https://www.googleapis.com/auth/drive.file'
TOKEN_FILE = '/tmp/drive_token.json'
REDIRECT_URI = 'https://lucaz-server.onrender.com/auth/callback'

STATUS_FOLDER_IDS = {
    'Active': '1qcWcpTDiY6gQDJlDhr76cNh9R5qD38dG',
    'Completed': '1gqxIiZN7i8ts-D0b0B98INm6sxa8vWTp',
    'Rejected': '14JZv7q4lRk2I2A-5tk3FEwJI5-beCXjx'
}

# ── OAUTH DRIVE ──
def get_drive_token():
    """Get valid access token, refreshing if needed"""
    # First try env var for refresh token (persists across deploys)
    refresh_token = os.environ.get('GOOGLE_REFRESH_TOKEN')
    
    # Then try token file
    if not refresh_token and os.path.exists(TOKEN_FILE):
        with open(TOKEN_FILE) as f:
            token_data = json.load(f)
        refresh_token = token_data.get('refresh_token')
    
    if not refresh_token:
        return None
    
    res = http_requests.post('https://oauth2.googleapis.com/token', data={
        'client_id': CLIENT_ID,
        'client_secret': CLIENT_SECRET,
        'refresh_token': refresh_token,
        'grant_type': 'refresh_token'
    })
    data = res.json()
    if 'access_token' in data:
        return data['access_token']
    print(f'Token refresh error: {data}', flush=True)
    return None

def get_or_create_folder(token, name, parent_id):
    """Find or create folder inside parent"""
    try:
        q = f"name='{name}' and '{parent_id}' in parents and mimeType='application/vnd.google-apps.folder' and trashed=false"
        res = http_requests.get(
            'https://www.googleapis.com/drive/v3/files',
            headers={'Authorization': f'Bearer {token}'},
            params={'q': q, 'fields': 'files(id,name)'}
        )
        data = res.json()
        if data.get('files'):
            return data['files'][0]['id']
        # Create
        res = http_requests.post(
            'https://www.googleapis.com/drive/v3/files',
            headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
            json={'name': name, 'mimeType': 'application/vnd.google-apps.folder', 'parents': [parent_id]}
        )
        return res.json().get('id')
    except Exception as e:
        print(f'Folder error: {e}', flush=True)
        return None

def save_to_drive(doc_bytes, file_name, client_name, status='Active', existing_file_id=None, proposal_num=None, date_issued=None):
    """Save proposal to Drive under Status/Client/Job #XXXX - Date/"""
    try:
        token = get_drive_token()
        if not token:
            print('DRIVE: No token', flush=True)
            return None
        status_id = STATUS_FOLDER_IDS.get(status, STATUS_FOLDER_IDS['Active'])
        client_id = get_or_create_folder(token, client_name, status_id)
        print(f'DRIVE: client_folder={client_id}', flush=True)
        if not client_id:
            return None

        # Create job subfolder: "Job #0147 - 05-28-2026"
        job_folder_name = f"Job #{proposal_num or '----'} - {date_issued or 'Unknown'}"
        job_id = get_or_create_folder(token, job_folder_name, client_id)
        if not job_id:
            job_id = client_id  # fallback to client folder

        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        if existing_file_id:
            res = http_requests.patch(
                f'https://www.googleapis.com/upload/drive/v3/files/{existing_file_id}',
                headers={'Authorization': f'Bearer {token}', 'Content-Type': mimetype},
                params={'uploadType': 'media', 'fields': 'id'},
                data=doc_bytes
            )
            return res.json().get('id', existing_file_id)

        boundary = f'lucaz_{int(time.time())}'
        meta = json.dumps({'name': file_name, 'parents': [job_id]})
        body = (
            f'--{boundary}\r\nContent-Type: application/json; charset=UTF-8\r\n\r\n'
            f'{meta}\r\n--{boundary}\r\nContent-Type: {mimetype}\r\n\r\n'
        ).encode() + doc_bytes + f'\r\n--{boundary}--'.encode()
        res = http_requests.post(
            'https://www.googleapis.com/upload/drive/v3/files',
            headers={'Authorization': f'Bearer {token}', 'Content-Type': f'multipart/related; boundary={boundary}'},
            params={'uploadType': 'multipart', 'fields': 'id'},
            data=body
        )
        result = res.json()
        print(f'DRIVE UPLOAD: {result}', flush=True)
        # Store job folder ID too so photos go in same folder
        file_id = result.get('id')
        if file_id:
            # Return both file ID and job folder ID
            return {'fileId': file_id, 'jobFolderId': job_id}
        return None
    except Exception as e:
        print(f'DRIVE ERROR: {e}', flush=True)
        import traceback; traceback.print_exc()
        return None

def move_drive_file(file_id, old_status, new_status, client_name=None, job_folder_id=None):
    """Move job folder (or file) between status folders"""
    try:
        token = get_drive_token()
        if not token:
            return False
        old_status_id = STATUS_FOLDER_IDS.get(old_status)
        new_status_id = STATUS_FOLDER_IDS.get(new_status)
        if not old_status_id or not new_status_id:
            return False

        # If we have a job folder ID, move the entire job folder
        if job_folder_id:
            # Find or create client folder in new status
            if client_name:
                new_client_id = get_or_create_folder(token, client_name, new_status_id)
            else:
                new_client_id = new_status_id
            # Get current parent of job folder
            res = http_requests.get(
                f'https://www.googleapis.com/drive/v3/files/{job_folder_id}',
                headers={'Authorization': f'Bearer {token}'},
                params={'fields': 'id,parents'}
            )
            current_parents = res.json().get('parents', [])
            remove_parents = ','.join(current_parents) if current_parents else ''
            # Move job folder to new client folder
            move_res = http_requests.patch(
                f'https://www.googleapis.com/drive/v3/files/{job_folder_id}',
                headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
                params={'addParents': new_client_id, 'removeParents': remove_parents, 'fields': 'id'},
                json={}
            )
            success = 'id' in move_res.json()
            # Clean up empty client folders in old status
            if success and current_parents:
                for parent_id in current_parents:
                    if parent_id in [old_status_id, new_status_id]: continue
                    check = http_requests.get(
                        'https://www.googleapis.com/drive/v3/files',
                        headers={'Authorization': f'Bearer {token}'},
                        params={'q': f"'{parent_id}' in parents and trashed=false", 'fields': 'files(id)'}
                    )
                    if not check.json().get('files'):
                        http_requests.delete(
                            f'https://www.googleapis.com/drive/v3/files/{parent_id}',
                            headers={'Authorization': f'Bearer {token}'}
                        )
            return success

        # Fallback: move just the file
        res = http_requests.get(
            f'https://www.googleapis.com/drive/v3/files/{file_id}',
            headers={'Authorization': f'Bearer {token}'},
            params={'fields': 'id,name,parents'}
        )
        file_data = res.json()
        current_parents = file_data.get('parents', [])
        if client_name:
            new_client_id = get_or_create_folder(token, client_name, new_status_id)
        else:
            new_client_id = new_status_id
        remove_parents = ','.join(current_parents) if current_parents else old_status_id
        res = http_requests.patch(
            f'https://www.googleapis.com/drive/v3/files/{file_id}',
            headers={'Authorization': f'Bearer {token}', 'Content-Type': 'application/json'},
            params={'addParents': new_client_id, 'removeParents': remove_parents, 'fields': 'id'},
            json={}
        )
        return 'id' in res.json()
    except Exception as e:
        print(f'Move error: {e}', flush=True)
        import traceback; traceback.print_exc()
        return False
        return False

# ── OAUTH ENDPOINTS ──
@app.route('/auth', methods=['GET'])
def auth():
    """Redirect to Google OAuth"""
    params = {
        'client_id': CLIENT_ID,
        'redirect_uri': REDIRECT_URI,
        'response_type': 'code',
        'scope': SCOPES,
        'access_type': 'offline',
        'prompt': 'consent'
    }
    from urllib.parse import urlencode
    url = 'https://accounts.google.com/o/oauth2/v2/auth?' + urlencode(params)
    return redirect(url)

@app.route('/auth/callback', methods=['GET'])
def auth_callback():
    """Handle OAuth callback, save refresh token"""
    code = flask_request.args.get('code')
    if not code:
        return jsonify({'error': 'No code'}), 400
    res = http_requests.post('https://oauth2.googleapis.com/token', data={
        'code': code,
        'client_id': CLIENT_ID,
        'client_secret': CLIENT_SECRET,
        'redirect_uri': REDIRECT_URI,
        'grant_type': 'authorization_code'
    })
    data = res.json()
    if 'refresh_token' in data:
        with open(TOKEN_FILE, 'w') as f:
            json.dump(data, f)
        refresh_token = data['refresh_token']
        return f'''<html><body style="font-family:sans-serif;text-align:center;padding:60px;max-width:600px;margin:0 auto">
            <h2 style="color:#2d8a4e">✓ Google Drive connected!</h2>
            <p>To make this permanent (so you never need to re-auth after deploys), copy the token below and add it as a Render environment variable:</p>
            <p><strong>Key:</strong> <code>GOOGLE_REFRESH_TOKEN</code></p>
            <p><strong>Value:</strong></p>
            <textarea style="width:100%;height:80px;font-size:11px;padding:8px" onclick="this.select()">{refresh_token}</textarea>
            <p style="color:#888;font-size:12px">Go to Render → your service → Environment → Add that variable → Save → Redeploy once more. After that, Drive works forever with no re-auth needed.</p>
        </body></html>'''
    return jsonify({'error': 'No refresh token', 'data': data}), 400

@app.route('/auth/status', methods=['GET'])
def auth_status():
    token = get_drive_token()
    return jsonify({'connected': bool(token)})


# ── PROPOSAL GENERATION ──
def set_cell_text(cell, new_text, bold=None, italic=None):
    for para in cell.paragraphs:
        if not para.runs: continue
        first = para.runs[0]
        for run in para.runs[1:]: run.text = ""
        first.text = new_text
        if bold is not None: first.bold = bold
        if italic is not None: first.italic = italic
        return

def set_row_cell_text(row_el, col_idx, text, bold=None):
    cells = row_el.findall(qn('w:tc'))
    if col_idx >= len(cells): return
    for p in cells[col_idx].findall(qn('w:p')):
        runs = p.findall(qn('w:r'))
        if runs:
            r = runs[0]
            t = r.find(qn('w:t'))
            if t is None:
                t = OxmlElement('w:t'); r.append(t)
            t.text = text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            rpr = r.find(qn('w:rPr'))
            if rpr is not None:
                if bold is not None:
                    b = rpr.find(qn('w:b'))
                    if bold and b is None:
                        b = OxmlElement('w:b'); rpr.insert(0, b)
                    elif not bold and b is not None:
                        rpr.remove(b)
                i_el = rpr.find(qn('w:i'))
                if i_el is not None: rpr.remove(i_el)
            for extra_r in runs[1:]:
                t2 = extra_r.find(qn('w:t'))
                if t2 is not None: t2.text = ''
            break

def get_para_text(el, doc):
    try:
        from docx.text.paragraph import Paragraph
        return Paragraph(el, doc).text.strip()
    except: return ''

def rebuild_paint_table(tbl, surfaces):
    hdr_el = copy.deepcopy(tbl.rows[0]._element)
    for row in list(tbl.rows):
        tbl._element.remove(row._element)
    tbl._element.append(copy.deepcopy(hdr_el))
    for idx, sf in enumerate(surfaces):
        qty = sf.get('qty')
        try: qty_int = int(str(qty)) if qty else 0
        except: qty_int = 0
        nm = f"{sf['name']} × {qty_int}" if qty_int > 1 else sf['name']
        new_row = copy.deepcopy(hdr_el)
        fill = 'FFFFFF' if idx % 2 == 0 else 'FAF5F5'
        for tc in new_row.findall(qn('w:tc')):
            shd = tc.find(f'.//{qn("w:shd")}')
            if shd is not None:
                shd.set(qn('w:fill'), fill)
                shd.set(qn('w:color'), 'auto')
            # Reset text color to black for data rows
            for run in tc.findall(f'.//{qn("w:r")}'):
                rpr = run.find(qn('w:rPr'))
                if rpr is not None:
                    clr = rpr.find(qn('w:color'))
                    if clr is not None:
                        clr.set(qn('w:val'), '000000')
                    # Remove bold from data rows
                    b = rpr.find(qn('w:b'))
                    if b is not None: rpr.remove(b)
        tbl._element.append(new_row)
        vals = [nm, sf.get('paint',''), sf.get('sheen',''), sf.get('color','TBD'), f"{sf.get('pc',2)} / {sf.get('prc',0)}"]
        for ci, val in enumerate(vals):
            set_row_cell_text(new_row, ci, val, bold=(ci==0))

def remove_side_block(doc, side_label):
    search = side_label + ' of House'
    body = doc.element.body
    children = list(body)
    for i, child in enumerate(children):
        if child.tag.split('}')[-1] == 'p' and search in get_para_text(child, doc):
            to_remove = [child]
            j = i + 1
            while j < len(children) and len(to_remove) <= 3:
                nc = children[j]
                nc_tag = nc.tag.split('}')[-1]
                nc_text = get_para_text(nc, doc) if nc_tag == 'p' else ''
                if nc_text and any(x in nc_text for x in ['of House','PROJECT PHOTOS','WARRANTY','PAYMENT','COST','NEXT','ESTIMATE','Inspection']):
                    break
                to_remove.append(nc)
                j += 1
            for el in to_remove:
                try: body.remove(el)
                except: pass
            return

SIDE_TABLE_IDX = {'Front': 3, 'Left': 4, 'Right': 5, 'Back': 6}

def _insert_note_para(doc, tbl_el, notes, body=None, idx=None):
    """Insert an italic note paragraph after a table element"""
    note_para = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '60')
    pPr.append(sp); note_para.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    i_el = OxmlElement('w:i')
    clr = OxmlElement('w:color'); clr.set(qn('w:val'), '555555')
    rPr.append(i_el); rPr.append(clr); r.append(rPr)
    t = OxmlElement('w:t')
    t.text = f'Note: {notes}'
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t); note_para.append(r)
    if tbl_el is not None and body is None:
        tbl_el.addnext(note_para)
    elif body is not None and idx is not None:
        body.insert(idx, note_para)

def generate_proposal(E):
    doc = Document(TEMPLATE_PATH)
    sides_data = E.get('sides', [])
    active_labels = [s['label'] for s in sides_data]
    all_sides = ['Front', 'Left', 'Right', 'Back']

    # 1. Footer page number fix
    for sect in doc.sections:
        ftr = sect.footer
        for para in ftr.paragraphs:
            for run in para.runs:
                if run.text == 'Page':
                    run.text = 'Page '
                    parent = run._element.getparent()
                    idx = list(parent).index(run._element)
                    for field_el in reversed(_make_page_field(run._element)):
                        parent.insert(idx + 1, field_el)
                    break

    # 2. Proposal # / License / Date
    for para in doc.paragraphs:
        if 'Proposal #:' in para.text and 'License:' in para.text:
            for run in para.runs:
                run.text = run.text.replace('0135', E.get('proposalNum', '____'))
                run.text = run.text.replace('05/14/2026', E.get('dateIssued', ''))
            break

    # 3. Client info
    ct = doc.tables[0]
    addr = f"{E['client']['street']}, {E['client']['city']}, {E['client']['state']} {E['client']['zip']}"
    set_cell_text(ct.rows[0].cells[1], E.get('subject',''), italic=True)
    set_cell_text(ct.rows[1].cells[1], E['client']['name'], italic=True)
    set_cell_text(ct.rows[2].cells[1], addr, italic=True)
    set_cell_text(ct.rows[3].cells[1], E['client']['phone'], italic=True)
    set_cell_text(ct.rows[4].cells[1], E['client']['email'], italic=True)

    # 4. Power wash bullets
    pw_cell = doc.tables[1].rows[0].cells[0]
    bps = [p for p in pw_cell.paragraphs if p.text.strip() and 'Power Washing' not in p.text]
    pw_items = [x for x in E.get('powerWash', []) if x.strip()]
    for i, item in enumerate(pw_items):
        if i < len(bps) and bps[i].runs:
            for r in bps[i].runs: r.text = ''
            bps[i].runs[0].text = item
    # Remove unused bullet paragraphs completely
    for i in range(len(pw_items), len(bps)):
        try: bps[i]._element.getparent().remove(bps[i]._element)
        except: pass

    # 5. Surface prep bullets
    sp_cell = doc.tables[2].rows[0].cells[0]
    sps = [p for p in sp_cell.paragraphs if p.text.strip() and 'Surface Preparation' not in p.text]
    sp_items = [x for x in E.get('surfacePrep', []) if x.strip()]
    for i, item in enumerate(sp_items):
        if i < len(sps) and sps[i].runs:
            for r in sps[i].runs: r.text = ''
            sps[i].runs[0].text = item
    for i in range(len(sp_items), len(sps)):
        try: sps[i]._element.getparent().remove(sps[i]._element)
        except: pass

    # 6. Rebuild paint tables
    for side in sides_data:
        label = side['label']
        tbl_idx = SIDE_TABLE_IDX.get(label)
        if tbl_idx is not None and tbl_idx < len(doc.tables):
            rebuild_paint_table(doc.tables[tbl_idx], side.get('surfaces', []))

    # 7. Remove unused sides
    for sl in [s for s in all_sides if s not in active_labels]:
        remove_side_block(doc, sl)

    # 8. Renumber side headings, insert custom sides BEFORE carpentry, add notes
    default_sides_list = ['Front', 'Left', 'Right', 'Back']
    # Template order for default sides
    template_order = ['Front', 'Left', 'Right', 'Back']
    
    # Find insertion point for custom sides — BEFORE carpentry section
    body = doc.element.body
    insert_before_el = None
    for child in list(body):
        if child.tag.split('}')[-1] == 'p':
            txt = get_para_text(child, doc)
            if 'Inspection & Carpentry' in txt or 'PROJECT PHOTOS' in txt:
                insert_before_el = child
                break

    # Renumber default sides based on their position in sides_data
    default_num = 3
    for side in sides_data:
        label = side['label']
        if label in default_sides_list:
            for para in doc.paragraphs:
                if label + ' of House' in para.text and para.runs:
                    para.runs[0].text = f"{default_num}.  {label} of House"
                    for r in para.runs[1:]: r.text = ''
                    break
            # Add note after paint table
            notes = side.get('notes', '').strip()
            if notes:
                tbl_idx = SIDE_TABLE_IDX.get(label)
                if tbl_idx is not None and tbl_idx < len(doc.tables):
                    _insert_note_para(doc, doc.tables[tbl_idx]._element, notes)
            default_num += 1

    # Insert custom sides before carpentry
    custom_num = default_num
    for side in sides_data:
        label = side['label']
        if label in default_sides_list:
            continue
        notes = side.get('notes', '').strip()

        if insert_before_el is not None:
            idx = list(body).index(insert_before_el)
        else:
            idx = len(list(body))

        # Heading paragraph — copy style from existing side heading
        heading = OxmlElement('w:p')
        import copy as _copy
        for para in doc.paragraphs:
            if 'of House' in para.text:
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None:
                    heading.append(_copy.deepcopy(pPr))
                if para.runs:
                    r = OxmlElement('w:r')
                    rPr = para.runs[0]._element.find(qn('w:rPr'))
                    if rPr is not None:
                        r.append(_copy.deepcopy(rPr))
                    t = OxmlElement('w:t')
                    t.text = f"{custom_num}.  {label}"
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    r.append(t); heading.append(r)
                break
        body.insert(idx, heading); idx += 1

        # Paint table
        if len(doc.tables) > 3:
            tbl_copy = _copy.deepcopy(doc.tables[3]._element)
            from docx.table import Table as DocxTable
            new_tbl = DocxTable(tbl_copy, doc)
            rebuild_paint_table(new_tbl, side.get('surfaces', []))
            body.insert(idx, tbl_copy); idx += 1
            if notes:
                _insert_note_para(doc, None, notes, body, idx)
                idx += 1

        # Spacer
        spacer = OxmlElement('w:p')
        body.insert(idx, spacer)
        insert_before_el = spacer
        custom_num += 1

    # Update carpentry number
    if E.get('carpentry', {}).get('enabled', False):
        carp_num = custom_num
        for para in doc.paragraphs:
            if 'Inspection & Carpentry' in para.text and para.runs:
                para.runs[0].text = f"{carp_num}.  Inspection & Carpentry (Work Change Order)"
                for r in para.runs[1:]: r.text = ''
                break

    # 9. Carpentry
    if not E.get('carpentry', {}).get('enabled', False):
        body = doc.element.body
        to_remove = []
        for child in list(body):
            tag = child.tag.split('}')[-1]
            if tag == 'p' and 'Inspection & Carpentry' in get_para_text(child, doc):
                to_remove.append(child)
            elif tag == 'tbl':
                try:
                    from docx.table import Table
                    if 'IMPORTANT' in Table(child, doc).rows[0].cells[0].text:
                        to_remove.append(child)
                except: pass
        for el in to_remove:
            try: body.remove(el)
            except: pass
    else:
        pass  # Carpentry numbering handled in step 8

    # 10. Duration
    duration = E.get('duration', '')
    if duration:
        for para in doc.paragraphs:
            if 'anticipated to take approximately' in para.text and '(X' in para.text:
                full = para.text.replace('(X\u2013X)', duration).replace('(X-X)', duration)
                if para.runs:
                    para.runs[0].text = full
                    for r in para.runs[1:]: r.text = ''
                break

    # 11. Cost table
    tax_pct = E.get('salesTaxPct', '8.375%')
    cost_map = {
        'Subtotal': E.get('subtotal',''),
        'Sales Tax': E.get('salesTaxAmt',''),
        'Total Cost for Project': E.get('total',''),
        'Deposit Due': E.get('deposit',''),
        'Balance Due': E.get('balance',''),
        'Porta Potty': E.get('portaPottyAmt','$200.00'),
    }
    for tbl in doc.tables:
        for row in tbl.rows:
            if len(row.cells) < 2: continue
            cell0 = row.cells[0].text.strip()
            # Update Sales Tax label with correct percentage
            if 'Sales Tax' in cell0:
                set_cell_text(row.cells[0], f'Sales Tax ({tax_pct})', italic=False)
                set_cell_text(row.cells[1], E.get('salesTaxAmt',''))
                continue
            for key, val in cost_map.items():
                if val and key in cell0:
                    set_cell_text(row.cells[1], val)
                    break

    # 12. Fix spacing before photos
    for para in doc.paragraphs:
        if 'Photos of the areas' in para.text:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr'); para._element.insert(0, pPr)
            sp = pPr.find(qn('w:spacing'))
            if sp is None:
                sp = OxmlElement('w:spacing'); pPr.append(sp)
            sp.set(qn('w:after'), '40')
            break

    # 13. Photos — build dynamic grid
    import base64
    from docx.shared import Inches
    photos = E.get('photos', {})
    photo_labels = E.get('photoLabels', {})

    label_to_key = {
        'Back':'Back','Left':'Left','Garage':'add1',
        'Front':'Front','Right':'Right',
        'Additional 1':'add1','Additional 2':'add2','Additional 3':'add3'
    }
    for key, lbl in photo_labels.items():
        if lbl: label_to_key[lbl] = key

    # Collect active photos in order
    active_photos = []
    for side in sides_data:
        key = side['label']
        pd = photos.get(key)
        lbl = photo_labels.get(key) or f"{side['label']} side"
        if pd and (pd.startswith('data:image') or (len(pd) < 100 and not pd.startswith('data:'))):
            active_photos.append((key, lbl, pd))
    # Collect all additional photos — any key not already used as a side photo
    side_keys_used = {s['label'] for s in sides_data}
    for slot_key, pd in photos.items():
        if slot_key in side_keys_used:
            continue
        if not pd:
            continue
        if pd.startswith('data:image') or (len(pd) < 100 and not pd.startswith('data:')):
            lbl = photo_labels.get(slot_key) or slot_key.replace('_',' ').title()
            active_photos.append((slot_key, lbl, pd))

    # Remove existing photo tables
    body = doc.element.body
    photo_tbls = []
    for tbl in doc.tables:
        if any(any('[ Insert Photo Here ]' in p.text for p in cell.paragraphs)
               for row in tbl.rows for cell in row.cells):
            photo_tbls.append(tbl)

    insert_before_el = None
    for child in list(body):
        if child.tag.split('}')[-1] == 'p':
            txt = get_para_text(child, doc)
            if 'PROJECT INFORMATION' in txt:
                insert_before_el = child
                break

    for tbl in photo_tbls:
        try: body.remove(tbl._element)
        except: pass

    if not active_photos:
        # Remove PROJECT PHOTOS section
        children = list(body)
        for i, child in enumerate(children):
            if child.tag.split('}')[-1] == 'p' and 'PROJECT PHOTOS' in get_para_text(child, doc):
                to_rm = [child]
                j = i + 1
                while j < len(children) and len(to_rm) < 4:
                    nc = children[j]
                    nc_txt = get_para_text(nc, doc) if nc.tag.split('}')[-1] == 'p' else ''
                    if nc_txt and any(x in nc_txt for x in ['PROJECT INFORMATION','WARRANTY','PAYMENT']):
                        break
                    to_rm.append(nc)
                    j += 1
                for el in to_rm:
                    try: body.remove(el)
                    except: pass
                break
    else:
        W = 9360
        rows = [active_photos[i:i+3] for i in range(0, len(active_photos), 3)]

        def make_cell(width, key, lbl, pd, doc_ref):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(width)); tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top','left','bottom','right']:
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), 'CCCCCC')
                tcBorders.append(b)
            tcPr.append(tcBorders); tc.append(tcPr)
            # Add image via temp paragraph on actual doc
            tmp = doc_ref.add_paragraph()
            tmp.alignment = 1
            if pd:
                try:
                    img_bytes = get_photo_bytes(pd)
                    if img_bytes:
                        run = tmp.add_run()
                        run.add_picture(io.BytesIO(img_bytes), width=Inches(min(2.3, width/1440.0)))
                except Exception as e:
                    print(f'Photo error: {e}', flush=True)
            pPr = tmp._element.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr'); tmp._element.insert(0, pPr)
            sp = OxmlElement('w:spacing'); sp.set(qn('w:before'), '60'); sp.set(qn('w:after'), '40')
            pPr.append(sp)
            jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'center'); pPr.append(jc)
            p1 = tmp._element
            p1.getparent().remove(p1)
            tc.append(p1)
            # Label
            p2 = OxmlElement('w:p')
            pPr2 = OxmlElement('w:pPr')
            jc2 = OxmlElement('w:jc'); jc2.set(qn('w:val'), 'center')
            sp2 = OxmlElement('w:spacing'); sp2.set(qn('w:before'), '0'); sp2.set(qn('w:after'), '60')
            pPr2.append(jc2); pPr2.append(sp2); p2.append(pPr2)
            r2 = OxmlElement('w:r')
            rPr2 = OxmlElement('w:rPr')
            b2 = OxmlElement('w:b')
            fn2 = OxmlElement('w:rFonts'); fn2.set(qn('w:ascii'), 'Calibri')
            rPr2.append(b2); rPr2.append(fn2); r2.append(rPr2)
            t2 = OxmlElement('w:t'); t2.text = lbl
            t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r2.append(t2); p2.append(r2); tc.append(p2)
            return tc

        def make_table(row_photos, doc_ref):
            n = len(row_photos); col_w = W // n
            tbl = OxmlElement('w:tbl')
            tblPr = OxmlElement('w:tblPr')
            tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(W)); tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW); tbl.append(tblPr)
            tblGrid = OxmlElement('w:tblGrid')
            for _ in row_photos:
                gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(col_w)); tblGrid.append(gc)
            tbl.append(tblGrid)
            tr = OxmlElement('w:tr')
            for k, l, pd in row_photos:
                tr.append(make_cell(col_w, k, l, pd, doc_ref))
            tbl.append(tr)
            return tbl

        if insert_before_el is not None:
            idx = list(body).index(insert_before_el)
        else:
            idx = len(list(body))

        spacer = OxmlElement('w:p')
        sp_pPr = OxmlElement('w:pPr')
        sp_sp = OxmlElement('w:spacing'); sp_sp.set(qn('w:before'), '80'); sp_sp.set(qn('w:after'), '0')
        sp_pPr.append(sp_sp); spacer.append(sp_pPr)

        for ri, row_photos in enumerate(rows):
            body.insert(idx, make_table(row_photos, doc))
            idx += 1
            if ri < len(rows) - 1:
                body.insert(idx, copy.deepcopy(spacer)); idx += 1

    # 14. Porta Potty
    if not E.get('portaPotty', False):
        for tbl in doc.tables:
            for row in list(tbl.rows):
                if 'Porta Potty' in row.cells[0].text:
                    row._element.getparent().remove(row._element)
                    break

    # 15. Signature section
    sig_tbl = None
    for tbl in doc.tables:
        if len(tbl.rows) > 0 and len(tbl.rows[0].cells) >= 3:
            if 'Client' in tbl.rows[0].cells[0].text or 'Signature' in tbl.rows[0].cells[0].text:
                sig_tbl = tbl; break

    if sig_tbl:
        def make_sig_line(space_before=160):
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            pBdr = OxmlElement('w:pBdr')
            bot = OxmlElement('w:bottom')
            bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
            bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), '000000')
            pBdr.append(bot)
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), str(space_before)); sp.set(qn('w:after'), '40')
            pPr.append(pBdr); pPr.append(sp); p.append(pPr)
            r = OxmlElement('w:r')
            t = OxmlElement('w:t'); t.text = ' '; r.append(t); p.append(r)
            return p

        def make_sig_label(text, bold=False):
            p = OxmlElement('w:p')
            pPr = OxmlElement('w:pPr')
            sp = OxmlElement('w:spacing')
            sp.set(qn('w:before'), '40'); sp.set(qn('w:after'), '20')
            pPr.append(sp); p.append(pPr)
            r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            fn = OxmlElement('w:rFonts'); fn.set(qn('w:ascii'), 'Calibri'); fn.set(qn('w:hAnsi'), 'Calibri')
            sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '18')
            rPr.append(fn); rPr.append(sz)
            if bold:
                b = OxmlElement('w:b'); rPr.append(b)
            r.append(rPr)
            t = OxmlElement('w:t'); t.text = text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t); p.append(r)
            return p

        for col_idx, name_lbl, sig_lbl in [(0,'Client Name','Client Signature'),(2,'Contractor','Authorized Signature')]:
            if col_idx >= len(sig_tbl.rows[0].cells): continue
            cell = sig_tbl.rows[0].cells[col_idx]
            for p_el in list(cell._element.findall(qn('w:p'))):
                cell._element.remove(p_el)
            cell._element.append(make_sig_label(''))
            cell._element.append(make_sig_line(200))
            cell._element.append(make_sig_label(name_lbl, bold=True))
            cell._element.append(make_sig_line(320))
            cell._element.append(make_sig_label(sig_lbl))
            cell._element.append(make_sig_line(320))
            cell._element.append(make_sig_label('Date'))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def _make_page_field(ref_run_el):
    rpr = ref_run_el.find(qn('w:rPr'))
    import copy as _copy
    def make_r():
        r = OxmlElement('w:r')
        if rpr is not None: r.append(_copy.deepcopy(rpr))
        return r
    r1 = make_r()
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), 'begin'); r1.append(fc)
    r2 = make_r()
    it = OxmlElement('w:instrText'); it.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve'); it.text = ' PAGE '; r2.append(it)
    r3 = make_r()
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'end'); r3.append(fc2)
    return [r1, r2, r3]


def get_photo_bytes(photo_ref):
    """Get image bytes from either a base64 data URL or a Drive file ID"""
    if not photo_ref:
        return None
    if photo_ref.startswith('data:image'):
        try:
            _, b64 = photo_ref.split(',', 1)
            return base64.b64decode(b64)
        except:
            return None
    else:
        # It's a Drive file ID
        try:
            token = get_drive_token()
            if not token:
                return None
            res = http_requests.get(
                f'https://www.googleapis.com/drive/v3/files/{photo_ref}',
                headers={'Authorization': f'Bearer {token}'},
                params={'alt': 'media'}
            )
            if res.status_code == 200:
                return res.content
        except Exception as e:
            print(f'Photo download error: {e}', flush=True)
        return None

# ── ROUTES ──
@app.route('/generate', methods=['POST', 'OPTIONS'])
def generate():
    if flask_request.method == 'OPTIONS':
        return '', 200
    try:
        E = flask_request.get_json()
        if not E:
            return jsonify({'error': 'No data'}), 400
        buf = generate_proposal(E)
        doc_bytes = buf.read()
        client_name = E.get('client', {}).get('name', 'Client')
        client_safe = client_name.replace(' ', '_')
        date = E.get('dateIssued', '').replace('/', '-')
        filename = f"LUCAZProposal_{client_safe}_{date}.docx"
        status = E.get('jobStatus', 'Active')
        existing_file_id = E.get('driveFileId', None)
        proposal_num = E.get('proposalNum', None)
        drive_result = save_to_drive(doc_bytes, filename, client_name, status, existing_file_id, proposal_num, date)
        print(f'DRIVE SAVE: {drive_result}', flush=True)
        response = send_file(
            io.BytesIO(doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
        if drive_result:
            if isinstance(drive_result, dict):
                response.headers['X-Drive-File-Id'] = drive_result.get('fileId', '')
                response.headers['X-Job-Folder-Id'] = drive_result.get('jobFolderId', '')
            else:
                response.headers['X-Drive-File-Id'] = drive_result
        return response
    except Exception as e:
        import traceback
        tb = traceback.format_exc()
        print('GENERATE ERROR:\n' + tb, flush=True)
        return jsonify({'error': str(e), 'trace': tb}), 500

@app.route('/move', methods=['POST', 'OPTIONS'])
def move():
    if flask_request.method == 'OPTIONS':
        return '', 200
    try:
        data = flask_request.get_json()
        success = move_drive_file(data.get('fileId'), data.get('oldStatus'), data.get('newStatus'), data.get('clientName'), data.get('jobFolderId'))
        return jsonify({'success': success})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/test-generate', methods=['GET'])
def test_generate():
    try:
        E = {
            'proposalNum':'0001','dateIssued':'05/27/2026','subject':'Test',
            'client':{'name':'Test Client','street':'123 Main St','city':'Ossining','state':'NY','zip':'10562','phone':'(914) 555-0000','email':'test@test.com'},
            'powerWash':['Full house exterior power wash'],'surfacePrep':['Scraping and sanding'],
            'sides':[{'label':'Front','surfaces':[{'name':'Siding — Clapboard','qty':None,'paint':'Regal Select','sheen':'Flat','color':'White','pc':2,'prc':0}]}],
            'carpentry':{'enabled':False},'portaPotty':False,'duration':'5-7 days',
            'photos':{},'subtotal':'$1,000.00','salesTaxAmt':'$83.75',
            'total':'$1,083.75','deposit':'$361.25','balance':'$722.50'
        }
        buf = generate_proposal(E)
        return jsonify({'success': True, 'size': len(buf.read())})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e), 'trace': traceback.format_exc()})

@app.route('/test-drive', methods=['GET'])
def test_drive():
    try:
        token = get_drive_token()
        if not token:
            return jsonify({'success': False, 'error': 'Not authorized — visit /auth first'})
        res = http_requests.get(
            'https://www.googleapis.com/drive/v3/files',
            headers={'Authorization': f'Bearer {token}'},
            params={'q': f"'{STATUS_FOLDER_IDS['Active']}' in parents and trashed=false", 'fields': 'files(id,name)', 'pageSize': 1}
        )
        data = res.json()
        if data.get('error'):
            return jsonify({'success': False, 'error': data['error']})
        return jsonify({'success': True, 'message': 'Drive connected', 'active_folder': STATUS_FOLDER_IDS['Active']})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok', 'template': os.path.exists(TEMPLATE_PATH), 'drive_auth': os.path.exists(TOKEN_FILE)})

@app.route('/upload-photo', methods=['POST', 'OPTIONS'])
def upload_photo():
    """Upload a photo to Drive and return the file ID"""
    if flask_request.method == 'OPTIONS':
        return '', 200
    try:
        token = get_drive_token()
        if not token:
            return jsonify({'error': 'Not authorized'}), 401
        data = flask_request.get_json()
        if not data:
            return jsonify({'error': 'No data'}), 400
        photo_data = data.get('photoData')  # base64 data URL
        client_name = data.get('clientName', 'Unknown')
        slot_label = data.get('slotLabel', 'Photo')
        if not photo_data or not photo_data.startswith('data:image'):
            return jsonify({'error': 'Invalid photo data'}), 400
        # Decode image
        import base64
        header, b64 = photo_data.split(',', 1)
        img_bytes = base64.b64decode(b64)
        # Determine mime type
        mime = 'image/jpeg'
        if 'png' in header: mime = 'image/png'
        # Get or create job folder — use jobFolderId if provided, else create from proposal info
        job_folder_id = data.get('jobFolderId')
        if job_folder_id:
            photos_id = get_or_create_folder(token, 'Photos', job_folder_id)
        else:
            # Create job folder structure from proposal info
            proposal_num = data.get('proposalNum', '----')
            date_issued = data.get('dateIssued', 'Unknown')
            status_id = STATUS_FOLDER_IDS.get('Active')
            client_id = get_or_create_folder(token, client_name, status_id)
            job_folder_name = f"Job #{proposal_num} - {date_issued}"
            job_id = get_or_create_folder(token, job_folder_name, client_id)
            photos_id = get_or_create_folder(token, 'Photos', job_id or client_id)
        # Upload photo
        boundary = f'photo_{int(time.time())}'
        ext = 'jpg' if mime == 'image/jpeg' else 'png'
        file_name = f'{slot_label}_{int(time.time())}.{ext}'
        meta = json.dumps({'name': file_name, 'parents': [photos_id]})
        body = (
            f'--{boundary}\r\nContent-Type: application/json; charset=UTF-8\r\n\r\n'
            f'{meta}\r\n--{boundary}\r\nContent-Type: {mime}\r\n\r\n'
        ).encode() + img_bytes + f'\r\n--{boundary}--'.encode()
        res = http_requests.post(
            'https://www.googleapis.com/upload/drive/v3/files',
            headers={'Authorization': f'Bearer {token}', 'Content-Type': f'multipart/related; boundary={boundary}'},
            params={'uploadType': 'multipart', 'fields': 'id'},
            data=body
        )
        result = res.json()
        print(f'PHOTO UPLOAD: {result}', flush=True)
        file_id = result.get('id')
        if not file_id:
            return jsonify({'error': 'Upload failed', 'detail': result}), 500
        # Return job folder ID so dashboard can store it for future uploads
        returned_job_folder = job_folder_id or (job_id if not job_folder_id else None)
        return jsonify({'success': True, 'fileId': file_id, 'jobFolderId': returned_job_folder})
    except Exception as e:
        import traceback
        print('PHOTO UPLOAD ERROR:', traceback.format_exc(), flush=True)
        return jsonify({'error': str(e)}), 500

@app.route('/photo/<file_id>', methods=['GET'])
def get_photo(file_id):
    """Serve a photo from Drive"""
    try:
        token = get_drive_token()
        if not token:
            return jsonify({'error': 'Not authorized'}), 401
        res = http_requests.get(
            f'https://www.googleapis.com/drive/v3/files/{file_id}',
            headers={'Authorization': f'Bearer {token}'},
            params={'alt': 'media'},
            stream=True
        )
        if res.status_code != 200:
            return jsonify({'error': 'Photo not found'}), 404
        from flask import Response
        return Response(
            res.content,
            mimetype=res.headers.get('Content-Type', 'image/jpeg'),
            headers={'Cache-Control': 'max-age=3600', 'Access-Control-Allow-Origin': '*'}
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete-photo/<file_id>', methods=['DELETE', 'OPTIONS'])
def delete_photo(file_id):
    """Delete a photo from Drive"""
    if flask_request.method == 'OPTIONS':
        return '', 200
    try:
        token = get_drive_token()
        if not token:
            return jsonify({'error': 'Not authorized'}), 401
        http_requests.delete(
            f'https://www.googleapis.com/drive/v3/files/{file_id}',
            headers={'Authorization': f'Bearer {token}'}
        )
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
